"""Temporary, validated extraction for uploaded vacancy documents."""

import os
import re
import tempfile
from pathlib import Path
from uuid import uuid4
from zipfile import BadZipFile

import fitz
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from starlette.concurrency import run_in_threadpool
from starlette.datastructures import UploadFile

from app.ai.client import AIClient
from app.config import get_max_job_document_size_bytes
from app.schemas.job_import import JobImportDraft
from app.services.job_import_service import generate_job_import_draft

_CHUNK_SIZE = 64 * 1024
_PDF_SIGNATURE = b"%PDF-"
_ZIP_SIGNATURE = b"PK\x03\x04"
_ALLOWED_MIME_TYPES = {
    ".pdf": {"application/pdf"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/zip",
    },
    ".txt": {"text/plain"},
}
_NO_TEXT_MESSAGE = (
    "No readable text was found. Please upload a text-based PDF, DOCX, or TXT file."
)


class JobDocumentError(Exception):
    pass


class UnsupportedJobDocumentError(JobDocumentError):
    pass


class JobDocumentTooLargeError(JobDocumentError):
    pass


class InvalidJobDocumentError(JobDocumentError):
    pass


class NoReadableJobDocumentTextError(JobDocumentError):
    pass


class JobDocumentProcessingError(JobDocumentError):
    pass


def normalize_job_document_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\x00", "")
    normalized = "\n".join(line.rstrip(" \t") for line in normalized.split("\n"))
    return re.sub(r"\n{3,}", "\n\n", normalized).strip()


def _safe_filename(filename: str | None) -> tuple[str, str]:
    if not filename:
        raise UnsupportedJobDocumentError("A job document filename is required.")
    basename = Path(filename.replace("\\", "/")).name.replace("\x00", "").strip()
    suffix = Path(basename).suffix.lower()
    if not basename or suffix not in _ALLOWED_MIME_TYPES:
        raise UnsupportedJobDocumentError(
            "Unsupported file format. Upload a PDF, DOCX, or TXT file."
        )
    if len(basename) > 255:
        basename = f"{Path(basename).stem[:250 - len(suffix)]}{suffix}"
    return basename, suffix


def _validate_mime_type(suffix: str, content_type: str | None) -> None:
    if (content_type or "").lower() not in _ALLOWED_MIME_TYPES[suffix]:
        raise UnsupportedJobDocumentError(
            f"The uploaded {suffix[1:].upper()} file has an unsupported MIME type."
        )


def _validate_signature(suffix: str, first_chunk: bytes) -> None:
    if suffix == ".pdf" and not first_chunk.startswith(_PDF_SIGNATURE):
        raise InvalidJobDocumentError("Invalid PDF file. The PDF signature is missing.")
    if suffix == ".docx" and not first_chunk.startswith(_ZIP_SIGNATURE):
        raise InvalidJobDocumentError("Invalid DOCX file. The ZIP signature is missing.")
    if suffix == ".txt" and b"\x00" in first_chunk:
        raise InvalidJobDocumentError("Invalid TXT file. Binary content is not supported.")


async def _stage_upload(upload: UploadFile, directory: Path) -> tuple[Path, str]:
    _, suffix = _safe_filename(upload.filename)
    _validate_mime_type(suffix, upload.content_type)
    path = directory / f"{uuid4()}{suffix}"
    maximum_size = get_max_job_document_size_bytes()
    bytes_written = 0
    first_chunk = True
    try:
        with path.open("xb") as destination:
            while chunk := await upload.read(_CHUNK_SIZE):
                if first_chunk:
                    first_chunk = False
                    _validate_signature(suffix, chunk)
                bytes_written += len(chunk)
                if bytes_written > maximum_size:
                    raise JobDocumentTooLargeError(
                        "Job document exceeds the configured maximum size."
                    )
                destination.write(chunk)
        if bytes_written == 0:
            raise InvalidJobDocumentError("The uploaded job document is empty.")
    except JobDocumentError:
        path.unlink(missing_ok=True)
        raise
    except OSError as exc:
        path.unlink(missing_ok=True)
        raise JobDocumentProcessingError(
            "Unable to process the job document."
        ) from exc
    return path, suffix


def _extract_pdf(path: Path) -> str:
    try:
        document = fitz.open(path)
    except (fitz.FileDataError, fitz.EmptyFileError, RuntimeError) as exc:
        raise InvalidJobDocumentError("Invalid or unreadable PDF file.") from exc
    with document:
        if document.needs_pass:
            raise InvalidJobDocumentError("Encrypted PDF files are not supported.")
        try:
            text = "\n\n".join(
                document.load_page(index).get_text("text")
                for index in range(document.page_count)
            )
        except (RuntimeError, ValueError) as exc:
            raise InvalidJobDocumentError("Unable to extract text from this PDF.") from exc
    return normalize_job_document_text(text)


def _extract_docx(path: Path) -> str:
    try:
        document = Document(path)
    except (PackageNotFoundError, BadZipFile, KeyError, ValueError) as exc:
        raise InvalidJobDocumentError("Invalid or unreadable DOCX file.") from exc
    content = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            content.extend(cell.text for cell in row.cells)
    return normalize_job_document_text("\n".join(content))


def _extract_txt(path: Path) -> str:
    try:
        content = path.read_bytes()
    except OSError as exc:
        raise JobDocumentProcessingError(
            "Unable to process the job document."
        ) from exc
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        try:
            text = content.decode("cp1252")
        except UnicodeDecodeError as exc:
            raise InvalidJobDocumentError(
                "TXT files must contain readable UTF-8 text."
            ) from exc
    return normalize_job_document_text(text)


def extract_job_document_text(path: Path, suffix: str) -> str:
    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    else:
        text = _extract_txt(path)
    word_count = len(re.findall(r"\b\w+\b", text, flags=re.UNICODE))
    if len(text) < 20 or word_count < 3:
        raise NoReadableJobDocumentTextError(_NO_TEXT_MESSAGE)
    return text


async def import_job_document(
    upload: UploadFile,
    *,
    ai_client: AIClient,
) -> JobImportDraft:
    """Extract and analyze a document while guaranteeing temporary cleanup."""

    try:
        with tempfile.TemporaryDirectory(prefix="recruitment-job-import-") as temp_dir:
            path, suffix = await _stage_upload(upload, Path(temp_dir))
            text = await run_in_threadpool(extract_job_document_text, path, suffix)
            return await generate_job_import_draft(text, ai_client=ai_client)
    except JobDocumentError:
        raise
    except OSError as exc:
        raise JobDocumentProcessingError(
            "Unable to process the job document."
        ) from exc
